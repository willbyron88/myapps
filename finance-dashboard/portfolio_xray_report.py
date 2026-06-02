#!/usr/bin/env python3
"""
Portfolio X-Ray v6 - Will Power finance-dashboard workflow with OpenAI web-search analysis tab

Default workflow on Will's PC:
    python portfolio_xray_report_v6_openai_websearch.py

The script automatically scans:
    C:\\Users\\willb\\myapps\\finance-dashboard

for the newest account summary / positions file, the prompt DOCX/TXT/TXT, and an optional
holdings_cache.csv. It writes output to:
    C:\\Users\\willb\\myapps\\finance-dashboard\\output\\latest_portfolio_xray.html

You can still override any path:
    python portfolio_xray_report_v6_openai_websearch.py --project-dir "C:\\Users\\willb\\myapps\\finance-dashboard"
    python portfolio_xray_report_v6_openai_websearch.py --positions "Account Summary _ Charles Schwab.pdf"

What it does:
  - Parses Schwab account-summary PDFs using Market Value, not cost basis.
  - Reads CSV/XLSX position exports when available.
  - Keeps crypto funds as fund-level crypto exposure.
  - Looks through mutual funds/ETFs using holdings_cache.csv first, then yfinance.
  - Uses only each fund's top 20 holdings, then keeps the remaining fund value as
    a separate FUND_OTHER bucket so totals reconcile to roughly 100%.
  - Produces a self-contained two-tab HTML dashboard:
      1) Portfolio X-Ray
      2) Rotation Review guided by your quarterly review DOCX prompt

Recommended optional holdings_cache.csv columns:
    fund_ticker, holding_ticker, holding_name, weight_pct, sector, theme

Install:
    pip install pandas openpyxl python-docx pymupdf yfinance requests beautifulsoup4 lxml openai python-dotenv

Not financial advice. Verify holdings and classifications before making decisions.
"""
from __future__ import annotations

import argparse
import datetime as dt
import html
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import docx
except Exception:
    docx = None

try:
    import yfinance as yf
except Exception:
    yf = None

try:
    from dotenv import load_dotenv
except Exception:
    load_dotenv = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

DEFAULT_PROJECT_DIR = Path(os.environ.get("FINANCE_DASHBOARD_DIR", r"C:\Users\willb\myapps\finance-dashboard"))
TOP_HOLDINGS_PER_FUND = 20

CRYPTO_FUNDS = {
    "FETH": "Ethereum",
    "ETH": "Ethereum",
    "FBTC": "Bitcoin",
    "IBIT": "Bitcoin",
    "SOLT": "Solana",
    "GXLM": "Stellar",
    "GBTC": "Bitcoin",
    "ETHE": "Ethereum",
    # Direct Coinbase / spot crypto tokens
    "BTC": "Bitcoin",
    "SOL": "Solana",
    "XLM": "Stellar",
    "XRP": "XRP",
    "LINK": "Chainlink",
    "ONDO": "Ondo",
    "AKT": "Akash Network",
    "AAVE": "Aave",
    "TAO": "Bittensor",
    "SYRUP": "Maple Finance",
    "FET": "Fetch.ai / ASI",
    "AERO": "Aerodrome Finance",
    "PYTH": "Pyth Network",
    "FLOCK": "Flock",
    "POL": "Polygon",
    "TRUMP": "Official Trump",
    "MATIC": "Polygon",
    "DOT": "Polkadot",
    "ADA": "Cardano",
    "AVAX": "Avalanche",
    "DOGE": "Dogecoin",
    "SHIB": "Shiba Inu",
    "LTC": "Litecoin",
    "UNI": "Uniswap",
    "ATOM": "Cosmos",
    "NEAR": "NEAR Protocol",
    "APT": "Aptos",
    "ARB": "Arbitrum",
    "OP": "Optimism",
    "INJ": "Injective",
    "SUI": "Sui",
    "JUP": "Jupiter",
}

# Categories that signal private / illiquid investments regardless of ticker
PRIVATE_INVESTMENT_PLATFORMS = {"republic", "wefunder", "seedinvest", "startengine", "mainvest", "netcapital"}

# Categories that signal pension / retirement accounts
PENSION_CATEGORIES = {"pension", "retirement", "401k", "403b", "ira", "annuity"}

CASH_LIKE = {"FDRXX", "SPAXX", "VMFXX", "SWVXX", "CASH", "MONEYMARKET"}

THEME_MAP = {
    "NVDA": "AI / Semiconductors / Compute",
    "TSM": "AI / Semiconductors / Compute",
    "TSMC": "AI / Semiconductors / Compute",
    "2330.TW": "AI / Semiconductors / Compute",
    "MU": "AI / Semiconductors / Compute",
    "CRDO": "AI / Semiconductors / Compute",
    "AVGO": "AI / Semiconductors / Compute",
    "AMD": "AI / Semiconductors / Compute",
    "ASML": "AI / Semiconductors / Compute",
    "ASML.AS": "AI / Semiconductors / Compute",
    "000660.KS": "AI / Semiconductors / Compute",
    "005930.KS": "AI / Semiconductors / Compute",
    "PLTR": "AI Software / Data Infrastructure",
    "BBAI": "AI Software / Data Infrastructure",
    "AISP": "AI Software / Data Infrastructure",
    "BRCHF": "AI / Semiconductors / Compute",
    "META": "US Large-Cap Growth / AI",
    "MSFT": "US Large-Cap Growth / AI",
    "GOOGL": "US Large-Cap Growth / AI",
    "GOOG": "US Large-Cap Growth / AI",
    "AMZN": "US Large-Cap Growth / AI",
    "AAPL": "US Large-Cap Growth",
    "OKLO": "Uranium / Nuclear / Energy Transition",
    "NXE": "Uranium / Nuclear / Energy Transition",
    "SRUUF": "Uranium / Nuclear / Energy Transition",
    "PGEZF": "Uranium / Nuclear / Energy Transition",
    "ABSI": "Biotech / AI Drug Discovery",
    "RXRX": "Biotech / AI Drug Discovery",
    "RLAY": "Biotech / AI Drug Discovery",
    "NTLA": "Biotech / Gene Editing",
    "LTRN": "Biotech / AI Drug Discovery",
    "TXG": "Biotech / Tools",
    "ASTS": "Space / Satellite Connectivity",
    "JOBY": "Aviation / Mobility",
    "O": "Real Estate / REITs",
    "ARCC": "Income / BDCs",
    "HTGC": "Income / BDCs",
    "USERX": "Gold / Precious Metals",
    "CVX": "Energy",
    "SLB": "Energy Services",
    "FETH": "Crypto / Tokenization",
    "ETH": "Crypto / Tokenization",
    "FBTC": "Crypto / Tokenization",
    "IBIT": "Crypto / Tokenization",
    "SOLT": "Crypto / Tokenization",
    "GXLM": "Crypto / Tokenization",
}

FUND_ASSET_CLASS_HINTS = {
    "FXAIX": "US Broad Market / S&P 500",
    "FLCNX": "US Large-Cap Growth",
    "FPADX": "Emerging Markets",
    "JEMWX": "Emerging Markets",
    "FSPSX": "International Developed",
    "FSSNX": "Small Cap",
    "FDSCX": "Small Cap",
    "AASRX": "Small Cap Value",
    "FSMAX": "US Extended Market",
    "FLKSX": "US Mid/Small Value / Low-Priced Stock",
    "FSRNX": "Real Estate / REITs",
    "CSJZX": "Real Estate / REITs",
    "FTKFX": "Bonds",
    "FXNAX": "Bonds",
    "USERX": "Gold / Precious Metals",
}

BAD_TICKERS = {
    "N/A", "NA", "IRA", "MORE", "YTD", "TYPE", "TOTAL", "PRICE", "RATINGS", "REINVEST",
    "QUANTITY", "DISCLOSURES", "BACK", "TOP", "FIDELITY", "SCHWAB", "BROKERAGE",
    "RETIREMENT", "INDIVIDUAL", "INVESTMENT", "COST", "GAIN", "LOSS", "ACCOUNT", "VALUE",
    "TODAY", "TABLE", "VIEW", "ADD", "POSITIONS"
}


def log(msg: str) -> None:
    print(msg, flush=True)


def money_to_float(s) -> Optional[float]:
    if s is None:
        return None
    s = str(s).replace("\u00a0", " ").strip()
    if not s or s.lower() == "nan":
        return None
    neg = "(" in s and ")" in s
    s2 = re.sub(r"[^0-9.\-]", "", s)
    if s2 in {"", ".", "-"}:
        return None
    try:
        v = float(s2)
        return -abs(v) if neg else v
    except ValueError:
        return None


def fmt_money(x: float) -> str:
    return f"${x:,.0f}"


def fmt_pct(x: float) -> str:
    return f"{x:.2f}%"


def resolve_project_dir(cli_project_dir: Optional[str]) -> Path:
    if cli_project_dir:
        return Path(cli_project_dir).expanduser().resolve()
    if DEFAULT_PROJECT_DIR.exists():
        return DEFAULT_PROJECT_DIR
    return Path.cwd().resolve()


def newest_file(project_dir: Path, patterns: List[str]) -> Optional[Path]:
    files: List[Path] = []
    assets_dir = project_dir / "assets"
    search_dirs = [assets_dir, project_dir] if assets_dir.is_dir() else [project_dir]
    for d in search_dirs:
        for pat in patterns:
            files.extend(d.glob(pat))
    files = [p for p in files if p.is_file() and not p.name.startswith("~$")]
    if not files:
        return None
    return max(files, key=lambda p: p.stat().st_mtime)


def looks_like_real_prompt(path: Path) -> bool:
    """Return True only for files that are likely to be Will's actual review prompt.

    This intentionally excludes requirements_*.txt, logs, output HTML, and other
    project support files. The previous version allowed any *.txt and could pick
    requirements_portfolio_xray_v4.txt as the prompt, which made the Analysis tab
    useless.
    """
    name = path.name.lower()
    if path.name.startswith("~$") or not path.is_file():
        return False
    bad_tokens = [
        "requirements", "debug", "openai_analysis", "response", "request",
        "portfolio_xray", "latest_portfolio", "sample", "output", "history",
        "cache", "holdings_cache", "positions", "account summary", "all-accounts"
    ]
    if any(tok in name for tok in bad_tokens):
        return False
    if path.suffix.lower() not in {".docx", ".txt", ".md"}:
        return False
    good_tokens = ["quarterly", "rotation", "review", "prompt", "pasted text"]
    if any(tok in name for tok in good_tokens):
        return True
    return False


def score_prompt_candidate(path: Path) -> tuple:
    """Higher tuple wins. Prefer DOCX and prompt-like names over generic pasted text."""
    name = path.name.lower()
    score = 0
    if "quarterly" in name: score += 50
    if "rotation" in name: score += 50
    if "review" in name: score += 30
    if "prompt" in name: score += 40
    if "pasted text" in name: score += 20
    if path.suffix.lower() == ".docx": score += 10
    if path.suffix.lower() == ".txt": score += 5
    return (score, path.stat().st_mtime)


def discover_prompt_file(project_dir: Path) -> Optional[Path]:
    candidates: List[Path] = []
    assets_dir = project_dir / "assets"
    search_dirs = [project_dir] + ([assets_dir] if assets_dir.is_dir() else [])
    for d in search_dirs:
        for pat in ["*.docx", "*.txt", "*.md"]:
            candidates.extend(d.glob(pat))
    candidates = [p for p in candidates if looks_like_real_prompt(p)]
    if not candidates:
        return None
    return max(candidates, key=score_prompt_candidate)


def discover_inputs(project_dir: Path, positions: Optional[str], prompt_file: Optional[str], holdings_cache: Optional[str], out: Optional[str]) -> Dict[str, Optional[Path]]:
    pos_path = Path(positions).expanduser() if positions else None
    if pos_path and not pos_path.is_absolute():
        pos_path = project_dir / pos_path
    if not pos_path:
        # Prefer PDFs first even if a CSV/XLSX was saved more recently.
        # Schwab PDFs contain the full market-value summary needed for this workflow.
        for group in (["Account Summary*.pdf", "*Account*Summary*.pdf"],
                      ["All-Accounts-Positions*.csv", "*Positions*.csv", "*positions*.csv"],
                      ["*.xlsx"]):
            pos_path = newest_file(project_dir, list(group))
            if pos_path:
                break

    prompt_path = Path(prompt_file).expanduser() if prompt_file else None
    if prompt_path and not prompt_path.is_absolute():
        prompt_path = project_dir / prompt_path
    if not prompt_path:
        # Pick only true prompt/review files. Never use requirements_*.txt or support files.
        prompt_path = discover_prompt_file(project_dir)

    cache_path = Path(holdings_cache).expanduser() if holdings_cache else None
    if cache_path and not cache_path.is_absolute():
        cache_path = project_dir / cache_path
    if not cache_path:
        c = project_dir / "holdings_cache.csv"
        cache_path = c if c.exists() else None

    out_path = Path(out).expanduser() if out else None
    if out_path and not out_path.is_absolute():
        out_path = project_dir / out_path
    if not out_path:
        output_dir = project_dir / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        out_path = output_dir / "latest_portfolio_xray.html"

    return {"positions": pos_path, "prompt_file": prompt_path, "prompt_docx": prompt_path, "holdings_cache": cache_path, "out": out_path}


def read_prompt_file(path: Optional[Path]) -> str:
    if not path or not path.exists():
        return ""
    if not looks_like_real_prompt(path):
        return ""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        if docx is None:
            return ""
        d = docx.Document(str(path))
        parts = []
        for para in d.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in d.tables:
            for row in table.rows:
                vals = [c.text.strip() for c in row.cells if c.text.strip()]
                if vals:
                    parts.append(" | ".join(vals))
        return "\n".join(parts)
    # Last resort for pasted or plain-text files with unusual extensions.
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


def read_prompt_docx(path: Optional[Path]) -> str:
    # Backwards-compatible wrapper.
    return read_prompt_file(path)


def extract_pdf_text(path: Path) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF is required for PDF parsing. pip install pymupdf")
    doc = fitz.open(str(path))
    return "\n".join(page.get_text() for page in doc)


def _is_ticker_line(s: str) -> bool:
    raw = str(s).upper().strip()
    if not re.fullmatch(r"[A-Z][A-Z0-9.\-]{0,7}", raw):
        return False
    return raw not in BAD_TICKERS


def _clean_ticker(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9.\-]", "", str(s).upper().strip())


def _is_unsigned_money_line(s: str) -> bool:
    s = str(s).replace("\u00a0", " ").strip()
    return bool(re.fullmatch(r"\$[0-9,]+(?:\.[0-9]+)?", s))


def _extract_row_pct(row_lines: List[str]) -> Optional[float]:
    pcts = []
    for ln in row_lines:
        x = str(ln).strip()
        if re.fullmatch(r"[0-9]+(?:\.[0-9]+)?%", x):
            try:
                pcts.append(float(x.replace("%", "")))
            except Exception:
                pass
    return pcts[-1] if pcts else None


def _choose_market_value(row_lines: List[str], total_value: Optional[float]) -> Optional[float]:
    unsigned = []
    for ln in row_lines:
        if _is_unsigned_money_line(ln):
            v = money_to_float(ln)
            if v is not None:
                unsigned.append(v)
    if not unsigned:
        return None

    # Primary: first unsigned dollar amount after a marker showing we've passed
    # price / price-change. This is usually the Schwab Market Value column.
    marker_seen = False
    for ln in row_lines[1:]:
        x = str(ln).strip()
        if re.fullmatch(r"\([+\-]?[0-9.]+%\)", x) or x == "--":
            marker_seen = True
            continue
        if _is_unsigned_money_line(x):
            v = money_to_float(x)
            if v is not None and marker_seen:
                return v

    # Fallback: Schwab % of holdings, if present.
    pct = _extract_row_pct(row_lines)
    if pct is not None and total_value:
        expected = total_value * pct / 100.0
        return min(unsigned, key=lambda x: abs(x - expected))

    return unsigned[0]


def _row_boundary_line(ln: str) -> bool:
    low = str(ln).strip().lower()
    if not low:
        return False
    starts = (
        "etfs & closed end funds total", "mutual funds total", "equities total",
        "cash & money market total", "others total", "etfs & closed end funds",
        "mutual funds", "cash & money market", "cash & cash investments",
        "symbol", "disclosures", "investment and insurance", "principal amount invested",
        "today's date", "check the background", "the charles schwab corporation",
        "https://", "account summary |", "5/29/", "(0219"
    )
    return any(low.startswith(s) for s in starts) or low == "others"


def parse_schwab_pdf(path: Path) -> Tuple[pd.DataFrame, Optional[float]]:
    text = extract_pdf_text(path)
    total_value = None
    m = re.search(r"Total Value\s*\n\$([0-9,]+\.[0-9]{2})", text)
    if m:
        total_value = money_to_float(m.group(1))

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    records = []
    section: Optional[str] = None
    active = False
    i = 0
    while i < len(lines):
        ln = lines[i]
        low = ln.lower()
        if low.startswith("equities") and "total" not in low:
            section = "Stock"; active = False; i += 1; continue
        if low.startswith("etfs & closed end funds") and "total" not in low:
            section = "ETF"; active = False; i += 1; continue
        if low.startswith("mutual funds") and "total" not in low:
            section = "Mutual Fund"; active = False; i += 1; continue
        if low.startswith("cash & money market") and "total" not in low:
            section = "Cash"; active = False; i += 1; continue
        if low == "others":
            section = "Other"; active = False; i += 1; continue
        if "symbol" in low and "name" in low:
            active = section is not None
            i += 1
            continue
        if low.startswith(("equities total", "etfs & closed end funds total", "mutual funds total", "cash & money market total", "others total")):
            active = False
            i += 1
            continue

        if active and section and _is_ticker_line(ln):
            j = i + 1
            while j < len(lines):
                if _row_boundary_line(lines[j]):
                    break
                if j > i + 1 and _is_ticker_line(lines[j]):
                    break
                j += 1
            row = lines[i:j]
            ticker = _clean_ticker(row[0])
            name = ""
            for cand in row[1:7]:
                cand = cand.strip()
                if re.fullmatch(r"\d+", cand) or cand in {"--", "Ratings"}:
                    continue
                if "$" not in cand and "%" not in cand and not re.fullmatch(r"[0-9,.]+", cand):
                    name = cand
                    break
            value = _choose_market_value(row, total_value)
            pct = _extract_row_pct(row)
            if value is not None and value >= 0 and ticker not in BAD_TICKERS:
                records.append({
                    "ticker": ticker,
                    "name": name or ticker,
                    "value": float(value),
                    "category": section,
                    "source": str(path.name),
                    "pdf_pct_of_holdings": pct,
                })
            i = j
            continue
        i += 1

    # Explicit Schwab non-symbol buckets.
    m_other = re.search(r"Others Total\s*\n\$([0-9,]+\.[0-9]{2})", text)
    if m_other:
        other_val = money_to_float(m_other.group(1))
        if other_val and other_val > 0:
            records.append({"ticker": "OTHER", "name": "Other holdings", "value": float(other_val), "category": "Other", "source": str(path.name), "pdf_pct_of_holdings": None})
    m_cash = re.search(r"Cash & Cash Investments\s*\n(?:\d+\s*\n)?\$([0-9,]+\.[0-9]{2})", text)
    if m_cash:
        cash_val = money_to_float(m_cash.group(1))
        if cash_val and cash_val > 0:
            records.append({"ticker": "CASH", "name": "Cash & Cash Investments", "value": float(cash_val), "category": "Cash", "source": str(path.name), "pdf_pct_of_holdings": None})

    df = pd.DataFrame(records)
    if not df.empty:
        df = df[~df["ticker"].isin(BAD_TICKERS)]
        df = df.drop_duplicates(subset=["ticker", "category"], keep="first")
    return df, total_value


SUPPORTED_ASSET_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".csv", ".txt", ".json", ".docx"}

# Column name aliases used for auto-detection (all lowercase, stripped).
# The first match wins per field. Add aliases here to extend without touching file code.
TICKER_ALIASES  = {"symbol", "ticker", "stock", "cusip", "isin", "security id"}
NAME_ALIASES    = {"investment", "name", "description", "security", "holding", "asset name", "security name"}
VALUE_ALIASES   = {"market value", "marketvalue", "current value", "mkt value", "value", "amount", "total value", "fair value", "balance"}
CATEGORY_ALIASES = {"account", "category", "type", "asset class", "asset type", "asset", "account type", "acct type"}


def _load_column_map(assets_dir: Path) -> Dict[str, str]:
    """Load optional column_map.json from assets/ dir. Keys are field names, values are exact column headers."""
    p = assets_dir / "column_map.json"
    if not p.exists():
        p = assets_dir.parent / "column_map.json"
    if not p.exists():
        return {}
    try:
        with open(p, encoding="utf-8") as f:
            return {k.lower(): v for k, v in json.load(f).items()}
    except Exception:
        return {}


def _resolve_col(cols: Dict[str, str], aliases: set, override: Optional[str] = None) -> Optional[str]:
    """Return the actual DataFrame column name matching the alias set or an explicit override."""
    if override:
        # Exact match first, then case-insensitive
        if override in cols.values():
            return override
        return cols.get(override.lower().strip())
    for key, real in cols.items():
        if key in aliases or any(a in key for a in aliases):
            return real
    return None


def _docx_to_dataframe(path: Path) -> pd.DataFrame:
    """Extract the first table from a DOCX file as a DataFrame."""
    if docx is None:
        raise RuntimeError("python-docx is required for DOCX parsing. pip install python-docx")
    doc = docx.Document(str(path))
    if not doc.tables:
        raise ValueError(f"No tables found in {path.name}. DOCX positions files must contain a table.")
    table = doc.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        raise ValueError(f"Empty table in {path.name}.")
    headers = rows[0]
    return pd.DataFrame(rows[1:], columns=headers)


def read_positions_table(path: Path, column_map: Optional[Dict[str, str]] = None) -> Tuple[pd.DataFrame, Optional[float]]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_schwab_pdf(path)

    if ext in {".xlsx", ".xls"}:
        raw = pd.read_excel(path, sheet_name=None)
        df = pd.concat(raw.values(), ignore_index=True)
    elif ext in {".csv", ".txt"}:
        df = pd.read_csv(path)
    elif ext == ".json":
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        df = pd.DataFrame(data if isinstance(data, list) else data.get("positions", data.get("holdings", [])))
    elif ext == ".docx":
        df = _docx_to_dataframe(path)
    else:
        raise ValueError(f"Unsupported positions file: {path.name}")

    if column_map is None:
        column_map = {}

    cols = {str(c).lower().strip(): c for c in df.columns}
    ticker_col  = _resolve_col(cols, TICKER_ALIASES,   column_map.get("ticker") or column_map.get("symbol"))
    name_col    = _resolve_col(cols, NAME_ALIASES,     column_map.get("name"))
    value_col   = _resolve_col(cols, VALUE_ALIASES,    column_map.get("value") or column_map.get("market value"))
    category_col = _resolve_col(cols, CATEGORY_ALIASES, column_map.get("category") or column_map.get("account"))

    if ticker_col is None or value_col is None:
        raise ValueError(
            f"Could not identify required columns in {path.name}.\n"
            f"  Columns found: {list(cols.keys())}\n"
            f"  Need at least: a ticker/symbol column and a market value column.\n"
            f"  Fix: rename your columns to match known aliases, or add a column_map.json to assets/.\n"
            f"  See assets/README.md for details."
        )

    out = pd.DataFrame()
    out["ticker"] = df[ticker_col].astype(str).str.upper().str.strip().map(_clean_ticker)
    out["name"] = df[name_col].astype(str) if name_col else out["ticker"]
    out["value"] = df[value_col].apply(money_to_float)
    out["category"] = df[category_col].astype(str) if category_col else "Unknown"
    out["source"] = str(path.name)
    out = out.dropna(subset=["ticker", "value"])
    out = out[(out["value"] > 0) & (~out["ticker"].isin(BAD_TICKERS))]
    return out, float(out["value"].sum())


def read_all_asset_files(project_dir: Path) -> Tuple[pd.DataFrame, Optional[float], List[Path]]:
    """Read every supported file in assets/ (or project root if no assets/ dir), merge into one DataFrame."""
    assets_dir = project_dir / "assets"
    search_dir = assets_dir if assets_dir.is_dir() else project_dir

    column_map = _load_column_map(search_dir)
    if column_map:
        log(f"  Column map loaded: {column_map}")

    skip_names = {"readme.md", "readme.txt", "column_map.json"}
    all_files = [
        p for p in sorted(search_dir.iterdir())
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_ASSET_EXTENSIONS
        and not p.name.startswith("~$")
        and p.name.lower() not in skip_names
        and not looks_like_real_prompt(p)
    ]

    if not all_files:
        return pd.DataFrame(), None, []

    frames: List[pd.DataFrame] = []
    loaded: List[Path] = []
    for path in all_files:
        try:
            df, _ = read_positions_table(path, column_map=column_map)
            if not df.empty:
                frames.append(df)
                loaded.append(path)
                log(f"  Loaded {len(df)} rows from {path.name}")
        except Exception as e:
            log(f"  Skipped {path.name}: {e}")

    if not frames:
        return pd.DataFrame(), None, []

    combined = pd.concat(frames, ignore_index=True)
    total = float(combined["value"].sum())
    return combined, total, loaded


def normalize_holding_ticker(t: str) -> str:
    t = str(t).upper().strip()
    aliases = {
        "TAIWAN SEMICONDUCTOR MANUFACTURING CO LTD": "TSM",
        "TAIWAN SEMICONDUCTOR": "TSM",
        "NVIDIA CORP": "NVDA",
        "MICROSOFT CORP": "MSFT",
        "APPLE INC": "AAPL",
        "AMAZON.COM INC": "AMZN",
        "ALPHABET INC CLASS A": "GOOGL",
        "ALPHABET INC CLASS C": "GOOG",
        "META PLATFORMS INC CLASS A": "META",
        "BERKSHIRE HATHAWAY INC CLASS A": "BRK.A",
        "BERKSHIRE HATHAWAY INC CLASS B": "BRK.B",
    }
    return aliases.get(t, t)


def get_yfinance_holdings(fund_ticker: str) -> pd.DataFrame:
    if yf is None:
        return pd.DataFrame()
    try:
        tk = yf.Ticker(fund_ticker)
        candidates = []
        fd = getattr(tk, "funds_data", None)
        if fd is not None:
            for attr in ["top_holdings", "equity_holdings"]:
                try:
                    x = getattr(fd, attr)
                    if isinstance(x, pd.DataFrame) and not x.empty:
                        candidates.append(x.copy())
                except Exception:
                    pass
        for attr in ["top_holdings", "holdings"]:
            try:
                x = getattr(tk, attr)
                if isinstance(x, pd.DataFrame) and not x.empty:
                    candidates.append(x.copy())
            except Exception:
                pass
        if not candidates:
            return pd.DataFrame()
        df = candidates[0].reset_index()
        cols = {str(c).lower(): c for c in df.columns}
        name_col = next((cols[c] for c in cols if "name" in c or c == "holding"), None)
        ticker_col = next((cols[c] for c in cols if "symbol" in c or "ticker" in c), None)
        if ticker_col is None and "index" in df.columns:
            ticker_col = "index"
        weight_col = next((c for c in df.columns if "weight" in str(c).lower() or "percent" in str(c).lower() or "holdingpercent" in str(c).lower()), None)
        if weight_col is None:
            return pd.DataFrame()
        out = pd.DataFrame()
        out["holding_ticker"] = df[ticker_col].astype(str) if ticker_col else df[name_col].astype(str)
        out["holding_name"] = df[name_col].astype(str) if name_col else out["holding_ticker"]
        w = pd.to_numeric(df[weight_col], errors="coerce")
        if len(w.dropna()) and w.dropna().max() <= 1.0:
            w = w * 100.0
        out["weight_pct"] = w
        out["fund_ticker"] = fund_ticker.upper()
        out = out.dropna(subset=["weight_pct"])
        out["holding_ticker"] = out["holding_ticker"].map(normalize_holding_ticker)
        return out[["fund_ticker", "holding_ticker", "holding_name", "weight_pct"]]
    except Exception:
        return pd.DataFrame()


def read_holdings_cache(path: Optional[Path]) -> pd.DataFrame:
    if not path or not path.exists():
        return pd.DataFrame()
    df = pd.read_csv(path)
    required = {"fund_ticker", "holding_ticker", "weight_pct"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Holdings cache missing columns: {missing}")
    if "holding_name" not in df.columns:
        df["holding_name"] = df["holding_ticker"]
    if "sector" not in df.columns:
        df["sector"] = ""
    if "theme" not in df.columns:
        df["theme"] = ""
    df["fund_ticker"] = df["fund_ticker"].astype(str).str.upper().str.strip()
    df["holding_ticker"] = df["holding_ticker"].astype(str).str.upper().str.strip().map(normalize_holding_ticker)
    df["weight_pct"] = pd.to_numeric(df["weight_pct"], errors="coerce")
    return df.dropna(subset=["weight_pct"])


def classify_position(row: pd.Series) -> str:
    t = str(row["ticker"]).upper()
    cat = str(row.get("category", ""))
    if t in CRYPTO_FUNDS:
        return "Crypto Fund"
    if t in CASH_LIKE or "cash" in cat.lower():
        return "Cash"
    if t in FUND_ASSET_CLASS_HINTS or "mutual" in cat.lower():
        return "Mutual Fund"
    if "etf" in cat.lower() or "closed" in cat.lower():
        return "ETF"
    if "equity" in cat.lower() or "stock" in cat.lower():
        return "Stock"
    return "Stock"


def build_xray(positions: pd.DataFrame, holdings_cache: pd.DataFrame, total_value: Optional[float]) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, Dict]:
    positions = positions.copy()
    positions["ticker"] = positions["ticker"].astype(str).str.upper().str.strip()
    positions["asset_type"] = positions.apply(classify_position, axis=1)
    if total_value is None or not total_value:
        total_value = float(positions["value"].sum())

    exposure_records = []
    coverage_records = []

    for _, r in positions.iterrows():
        t = str(r["ticker"]).upper()
        v = float(r["value"])
        typ = r["asset_type"]

        if typ in {"Mutual Fund", "ETF"} and t not in CRYPTO_FUNDS:
            h = pd.DataFrame()
            if not holdings_cache.empty:
                h = holdings_cache[holdings_cache["fund_ticker"].str.upper() == t].copy()
            if h.empty:
                h = get_yfinance_holdings(t)
            if not h.empty:
                h["weight_pct"] = pd.to_numeric(h["weight_pct"], errors="coerce")
                h = h.dropna(subset=["weight_pct"]).sort_values("weight_pct", ascending=False).head(TOP_HOLDINGS_PER_FUND)
                looked_pct = max(0.0, min(100.0, float(h["weight_pct"].sum())))
                for _, hr in h.iterrows():
                    w = max(0.0, min(100.0, float(hr["weight_pct"])))
                    ev = v * w / 100.0
                    ht = str(hr["holding_ticker"]).upper()
                    exposure_records.append({
                        "exposure_ticker": ht,
                        "exposure_name": str(hr.get("holding_name", ht)),
                        "value": ev,
                        "pct_portfolio": ev / total_value * 100,
                        "source": t,
                        "source_name": r.get("name", t),
                        "source_value": v,
                        "source_weight_pct": w,
                        "theme": str(hr.get("theme", "")).strip() or THEME_MAP.get(ht, "Look-through Equity"),
                        "sector": str(hr.get("sector", "")).strip() or "Unknown",
                        "is_stock_exposure": True,
                        "is_remainder": False,
                    })
                rem = max(0.0, 100.0 - looked_pct)
                if rem > 0.01:
                    exposure_records.append({
                        "exposure_ticker": f"{t}_OTHER",
                        "exposure_name": f"{t} other holdings outside top {TOP_HOLDINGS_PER_FUND}",
                        "value": v * rem / 100.0,
                        "pct_portfolio": v * rem / 100.0 / total_value * 100,
                        "source": t,
                        "source_name": r.get("name", t),
                        "source_value": v,
                        "source_weight_pct": rem,
                        "theme": FUND_ASSET_CLASS_HINTS.get(t, "Fund Remainder"),
                        "sector": "Diversified / Other",
                        "is_stock_exposure": False,
                        "is_remainder": True,
                    })
                coverage_records.append({"fund": t, "fund_name": r.get("name", t), "fund_value": v, "looked_through_pct": looked_pct, "remainder_pct": max(0.0, 100.0 - looked_pct), "source": "holdings_cache/yfinance"})
            else:
                exposure_records.append({
                    "exposure_ticker": f"{t}_OTHER",
                    "exposure_name": f"{t} fund holdings not available",
                    "value": v,
                    "pct_portfolio": v / total_value * 100,
                    "source": t,
                    "source_name": r.get("name", t),
                    "source_value": v,
                    "source_weight_pct": 100.0,
                    "theme": FUND_ASSET_CLASS_HINTS.get(t, "Fund - no look-through"),
                    "sector": "Fund",
                    "is_stock_exposure": False,
                    "is_remainder": True,
                })
                coverage_records.append({"fund": t, "fund_name": r.get("name", t), "fund_value": v, "looked_through_pct": 0.0, "remainder_pct": 100.0, "source": "not available"})
        else:
            name = CRYPTO_FUNDS.get(t, r.get("name", t)) if t in CRYPTO_FUNDS else r.get("name", t)
            category_lower = str(r.get("category", "")).lower().strip()
            # Derive theme: explicit map first, then category-based inference, then fallback
            if t in THEME_MAP:
                theme = THEME_MAP[t]
            elif t in CRYPTO_FUNDS:
                theme = "Crypto / Tokenization"
            elif any(p in category_lower for p in PRIVATE_INVESTMENT_PLATFORMS):
                theme = "Private / Startup Investments"
            elif any(p in category_lower for p in PENSION_CATEGORIES):
                theme = "Pension / Retirement Plan"
            elif "coinbase" in category_lower or "crypto" in category_lower or "wallet" in category_lower:
                theme = "Crypto / Tokenization"
            elif typ == "Cash":
                theme = "Cash"
            else:
                theme = "Direct Stock / Other"

            sector = "Crypto" if (t in CRYPTO_FUNDS or "coinbase" in category_lower or "crypto" in category_lower) else "Unknown"
            exposure_records.append({
                "exposure_ticker": t,
                "exposure_name": name,
                "value": v,
                "pct_portfolio": v / total_value * 100,
                "source": "direct",
                "source_name": "Direct holding",
                "source_value": v,
                "source_weight_pct": 100.0,
                "theme": theme,
                "sector": sector,
                "is_stock_exposure": typ == "Stock" and theme not in {"Private / Startup Investments", "Pension / Retirement Plan", "Crypto / Tokenization"},
                "is_remainder": False,
            })

    exp = pd.DataFrame(exposure_records)
    all_agg = (exp.groupby(["exposure_ticker", "exposure_name", "is_stock_exposure", "is_remainder"], as_index=False)
                 .agg(value=("value", "sum"), pct_portfolio=("pct_portfolio", "sum")))
    all_agg = all_agg.sort_values("value", ascending=False)

    stock_agg = all_agg[(all_agg["is_stock_exposure"] == True) & (all_agg["is_remainder"] == False)].copy()
    stock_agg = stock_agg.sort_values("value", ascending=False)

    detail = exp.sort_values(["exposure_ticker", "value"], ascending=[True, False])

    buckets = (exp.groupby("theme", as_index=False)
                 .agg(value=("value", "sum"), pct_portfolio=("pct_portfolio", "sum")))
    buckets = buckets.sort_values("value", ascending=False)

    meta = {
        "total_value": float(total_value),
        "generated_at": dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "fund_coverage": coverage_records,
        "top_holdings_per_fund": TOP_HOLDINGS_PER_FUND,
    }
    return all_agg, stock_agg, detail, buckets, meta


def append_total_row(df: pd.DataFrame, label_col: str, total_value: float, columns_to_sum: List[str]) -> pd.DataFrame:
    if df.empty:
        return df
    row = {c: "" for c in df.columns}
    row[label_col] = "TOTAL SHOWN"
    for c in columns_to_sum:
        if c in df.columns:
            row[c] = float(pd.to_numeric(df[c], errors="coerce").fillna(0).sum())
    if "pct_portfolio" in df.columns:
        row["pct_portfolio"] = float(pd.to_numeric(df["pct_portfolio"], errors="coerce").fillna(0).sum())
    return pd.concat([df, pd.DataFrame([row])], ignore_index=True)


def df_to_html_table(df: pd.DataFrame, columns: List[str], money_cols=(), pct_cols=(), max_rows=None, table_id="", sortable=True) -> str:
    use = df.head(max_rows).copy() if max_rows else df.copy()
    rows = []
    sort_attr = f" onclick=\"sortTable('{table_id}',this)\" style='cursor:pointer;user-select:none'" if sortable else ""
    header = "".join(f"<th data-col='{i}'{sort_attr}>{html.escape(str(c))} <span class='sort-icon'>&#8597;</span></th>" for i, c in enumerate(columns))
    rows.append(f"<thead><tr>{header}</tr></thead><tbody>")
    for _, r in use.iterrows():
        is_total = str(r.get(columns[0] if columns else "", "")).startswith("TOTAL")
        row_class = " class='total-row'" if is_total else ""
        cells = []
        for c in columns:
            val = r.get(c, "")
            raw_num = None
            if c in money_cols and val != "":
                try:
                    raw_num = float(val)
                    val = fmt_money(raw_num)
                except Exception:
                    pass
            elif c in pct_cols and val != "":
                try:
                    raw_num = float(val)
                    val = fmt_pct(raw_num)
                except Exception:
                    pass
            data_attr = f" data-val='{raw_num}'" if raw_num is not None else ""
            cells.append(f"<td{data_attr}>{html.escape(str(val))}</td>")
        rows.append(f"<tr{row_class}>" + "".join(cells) + "</tr>")
    rows.append("</tbody>")
    return f"<table id='{table_id}'>" + "\n".join(rows) + "</table>"



def build_portfolio_summary_for_llm(
    positions: pd.DataFrame,
    all_agg: pd.DataFrame,
    stock_agg: pd.DataFrame,
    detail: pd.DataFrame,
    buckets: pd.DataFrame,
    meta: Dict,
) -> Dict:
    """Build a compact, serializable portfolio payload for the OpenAI analysis tab."""
    total = float(meta.get("total_value", 0) or 0)

    def records(df: pd.DataFrame, cols: List[str], n: int = 25) -> List[Dict]:
        if df is None or df.empty:
            return []
        out = []
        for _, r in df.head(n).iterrows():
            item = {}
            for c in cols:
                if c in df.columns:
                    v = r.get(c)
                    if pd.isna(v):
                        v = None
                    elif isinstance(v, (pd.Timestamp, dt.datetime, dt.date)):
                        v = str(v)
                    elif isinstance(v, (float, int)):
                        v = float(v)
                    item[c] = v
            out.append(item)
        return out

    crypto_value = float(buckets[buckets["theme"].str.contains("Crypto", case=False, na=False)]["value"].sum()) if not buckets.empty else 0.0
    ai_value = float(buckets[buckets["theme"].str.contains("AI|Semiconductors|Compute", case=False, na=False)]["value"].sum()) if not buckets.empty else 0.0
    fund_remainder_value = float(detail[detail["is_remainder"] == True]["value"].sum()) if not detail.empty and "is_remainder" in detail.columns else 0.0

    positions_view = positions.copy()
    if total > 0 and "pct_portfolio" not in positions_view.columns:
        positions_view["pct_portfolio"] = positions_view["value"] / total * 100
    positions_view = positions_view.sort_values("value", ascending=False)

    detail_view = detail.copy()
    if not detail_view.empty:
        detail_view = detail_view.sort_values("value", ascending=False)

    payload = {
        "generated_at": meta.get("generated_at"),
        "total_value": total,
        "portfolio_guardrails": {
            "crypto_target_max_pct": "roughly 10% to 15%",
            "private_startup_illiquid_max_pct": "roughly 2% to 3%",
            "single_speculative_stock_watch_pct": "about 2%",
            "broad_core_retirement_mutual_fund_base_preference": "above 70% if possible",
            "preference": "Use future contributions before selling existing winners unless an exposure is clearly oversized or thesis has changed.",
        },
        "summary_metrics": {
            "crypto_value": crypto_value,
            "crypto_pct": crypto_value / total * 100 if total else 0,
            "ai_semis_compute_value": ai_value,
            "ai_semis_compute_pct": ai_value / total * 100 if total else 0,
            "fund_other_unmapped_value": fund_remainder_value,
            "fund_other_unmapped_pct": fund_remainder_value / total * 100 if total else 0,
        },
        "top_50_true_stock_concentrations": records(
            stock_agg,
            ["exposure_ticker", "exposure_name", "value", "pct_portfolio", "theme"],
            50,
        ),
        "all_exposures_including_fund_other": records(
            all_agg,
            ["exposure_ticker", "exposure_name", "value", "pct_portfolio", "theme", "is_remainder"],
            70,
        ),
        "theme_buckets": records(buckets, ["theme", "value", "pct_portfolio"], 40),
        "largest_original_positions": records(positions_view, ["ticker", "name", "asset_type", "category", "value", "pct_portfolio"], 60),
        "fund_lookthrough_coverage": meta.get("fund_coverage", []),
        "underlying_source_detail_top_records": records(
            detail_view,
            ["exposure_ticker", "exposure_name", "source", "source_name", "source_weight_pct", "value", "pct_portfolio", "theme", "is_remainder"],
            100,
        ),
        "important_method_notes": [
            "Direct crypto ETFs/funds are kept as fund-level crypto exposure rather than decomposed into coins unless explicitly mapped.",
            "Mutual fund and ETF look-through uses top holdings available from holdings_cache.csv or yfinance; remaining fund value is retained as a separate OTHER bucket.",
            "For the Rotation Review tab, the script uses the OpenAI Responses API with web_search enabled for current institutional outlooks; do not invent citations if a source is not found.",
            "This is not financial advice; the output is for decision support and sizing awareness.",
        ],
    }
    return payload


def markdownish_to_html(text: str) -> str:
    """Small markdown-to-HTML converter for OpenAI narrative output without extra dependencies."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n")
    lines = text.split("\n")
    out = []
    in_ul = False
    in_ol = False
    in_pre = False

    def close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            close_lists()
            if not in_pre:
                out.append("<pre style='white-space:pre-wrap;color:#A8B3C6'>")
                in_pre = True
            else:
                out.append("</pre>")
                in_pre = False
            continue
        if in_pre:
            out.append(html.escape(line))
            continue
        if not stripped:
            close_lists()
            continue
        if stripped.startswith("### "):
            close_lists(); out.append(f"<h3>{html.escape(stripped[4:])}</h3>"); continue
        if stripped.startswith("## "):
            close_lists(); out.append(f"<h2>{html.escape(stripped[3:])}</h2>"); continue
        if stripped.startswith("# "):
            close_lists(); out.append(f"<h2>{html.escape(stripped[2:])}</h2>"); continue
        m = re.match(r"^[-*]\s+(.+)$", stripped)
        if m:
            if not in_ul:
                close_lists(); out.append("<ul>"); in_ul = True
            out.append(f"<li>{html.escape(m.group(1))}</li>")
            continue
        m = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if m:
            if not in_ol:
                close_lists(); out.append("<ol>"); in_ol = True
            out.append(f"<li>{html.escape(m.group(1))}</li>")
            continue
        close_lists()
        safe = html.escape(stripped)
        safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
        out.append(f"<p>{safe}</p>")
    close_lists()
    if in_pre:
        out.append("</pre>")
    return "\n".join(out)


def generate_openai_narrative(
    positions: pd.DataFrame,
    all_agg: pd.DataFrame,
    stock_agg: pd.DataFrame,
    detail: pd.DataFrame,
    buckets: pd.DataFrame,
    meta: Dict,
    prompt_text: str,
    project_dir: Path,
) -> Tuple[Optional[str], str]:
    """Return (html_narrative, status).

    v6 behavior:
      - Loads OPENAI_API_KEY from the finance-dashboard .env file
      - Sends BOTH Will's prompt text and calculated portfolio JSON to OpenAI
      - Uses the OpenAI Responses API with the hosted web_search tool so the
        institutional-outlook section can find current sources/citations.
      - Writes exact debug payloads to output/debug/ so Will can verify what was sent.
      - Falls back to local rules-based analysis if the API call fails.
    """
    # Load .env from the finance-dashboard project folder first, then normal dotenv discovery.
    if load_dotenv is not None:
        env_path = project_dir / ".env"
        if env_path.exists():
            load_dotenv(dotenv_path=env_path, override=False)
        else:
            load_dotenv(override=False)

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None, "OPENAI_API_KEY not found; used local rules-based analysis."
    if OpenAI is None:
        return None, "openai package not installed; run: pip install -r requirements_portfolio_xray_v6.txt"

    # Use a web-search capable Responses model. Override in .env if desired.
    # Example: OPENAI_MODEL=gpt-5-mini or OPENAI_MODEL=gpt-5.5
    model = os.getenv("OPENAI_MODEL", "gpt-5-mini")
    log(f"OpenAI analysis: model={model}, web_search=on. This typically takes 1-3 minutes.")
    payload = build_portfolio_summary_for_llm(positions, all_agg, stock_agg, detail, buckets, meta)

    system_prompt = (
        "You are a cautious but growth-oriented portfolio analyst generating the SECOND TAB "
        "of a local HTML report for Will. The user's quarterly review prompt is the controlling "
        "specification. Treat the supplied calculated portfolio JSON as the factual source for "
        "Will's positions, percentages, look-through exposures, crypto exposure, AI exposure, "
        "theme buckets, and concentration risk. Use web search for current institutional outlooks. "
        "Cite institutional views with source names and URLs or source titles. Do not invent citations. "
        "Separate existing holdings from new-money/contribution guidance. Do not give directive personal "
        "financial advice; frame recommendations as decision-support options. The user prefers measured, "
        "contribution-first rebalancing and guardrails over drastic moves. Return Markdown-style text only; "
        "the Python script will convert it to HTML."
    )

    required_sections = [
        "1. What the Big Boys Are Saying (3–6 month and 12-month views from Goldman Sachs, JPMorgan, Morgan Stanley, BlackRock — cite each source)",
        "2. Where My Portfolio Agrees or Diverges (2–3 bullet points per major theme: crypto, AI/semis, broad core, international)",
        "3. Action Table (one row per significant position or theme: Hold / Add / Trim / Watch — with a one-line reason)",
        "4. New Money Allocation (where to direct next contributions — 3 to 5 bullet points, ranked by priority)",
        "5. Plain-English Bottom Line (3 to 5 sentences — what matters most this quarter)",
    ]

    web_instructions = """
IMPORTANT WEB SEARCH INSTRUCTIONS:
- You have web search available through the API. Use it.
- Search current institutional outlooks from BlackRock Investment Institute, Goldman Sachs, J.P. Morgan, Morgan Stanley, Fidelity, Vanguard, Bank of America, and State Street where available.
- Prefer 2026 outlooks, midyear outlooks, quarterly outlooks, capital market assumptions, or house-view pages.
- For each institution found, summarize: 3-6 month view, 12-month view, 3-5 year/longer view if available, and overweight/neutral/underweight views.
- If a specific institution's current outlook cannot be found, say 'not found in this run' rather than inventing it.
- Include source names/titles and links/citations in the text.
""".strip()

    user_content = (
        web_instructions
        + "\n\nQUARTERLY REVIEW PROMPT FROM FILE:\n"
        + (prompt_text[:50000] if prompt_text else "No quarterly prompt text found. Tell the user the prompt file was not detected.")
        + "\n\nCALCULATED PORTFOLIO X-RAY DATA JSON:\n"
        + json.dumps(payload, indent=2)
        + "\n\nMANDATORY OUTPUT REQUIREMENTS:\n"
        + "- Be concise. This is an executive summary, not a research report. Aim for quality over quantity.\n"
        + "- Use actual numbers from the JSON (crypto %, AI %, top holdings, theme buckets). Do not speak in generalities.\n"
        + "- Include exactly these sections in this order:\n" + "\n".join(required_sections) + "\n"
        + "- The Action Table must use real tickers or theme names from the JSON, not generic examples.\n"
        + "- Use plain language: Hold, Add, Trim, Watch. Skip jargon.\n"
        + "- Total length should be readable in under 5 minutes.\n"
    )

    # Write exact LLM request artifacts locally so you can inspect what the analysis tab used.
    debug_dir = project_dir / "output" / "debug"
    debug_dir.mkdir(parents=True, exist_ok=True)
    request_debug = {
        "api": "OpenAI Responses API",
        "model": model,
        "tools": [{"type": "web_search", "search_context_size": os.getenv("OPENAI_WEB_SEARCH_CONTEXT", "high")}],
        "system_prompt": system_prompt,
        "prompt_text_found": bool(prompt_text.strip()),
        "prompt_text_preview": prompt_text[:4000],
        "portfolio_payload": payload,
        "required_sections": required_sections,
    }
    (debug_dir / "openai_analysis_request.json").write_text(json.dumps(request_debug, indent=2), encoding="utf-8")
    (debug_dir / "openai_user_message.txt").write_text(user_content, encoding="utf-8")

    try:
        import threading, time as _time

        client = OpenAI(api_key=api_key)

        # Run the blocking API call in a background thread so we can show a spinner.
        _result: Dict = {}
        def _call():
            try:
                _result["resp"] = client.responses.create(
                    model=model,
                    tools=[{"type": "web_search", "search_context_size": os.getenv("OPENAI_WEB_SEARCH_CONTEXT", "high")}],
                    input=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_content},
                    ],
                )
            except Exception as e:
                _result["error"] = e

        t = threading.Thread(target=_call, daemon=True)
        t.start()

        phases = [
            (10,  "Sending portfolio data to OpenAI..."),
            (30,  "Searching institutional outlooks (Goldman, JPMorgan, BlackRock...)"),
            (60,  "Searching Morgan Stanley, Fidelity, Vanguard outlooks..."),
            (90,  "Analyzing portfolio vs institutional views..."),
            (120, "Generating rotation recommendations..."),
            (150, "Writing report sections..."),
            (999, "Finalizing analysis (this can take a few more minutes)..."),
        ]
        spin = ["|", "/", "-", "\\"]
        elapsed = 0
        phase_idx = 0
        phase_label = phases[0][1]
        print(f"  {phase_label}", end="", flush=True)
        while t.is_alive():
            _time.sleep(0.25)
            elapsed += 0.25
            s = spin[int(elapsed * 4) % 4]
            # Advance phase label when threshold is crossed
            while phase_idx < len(phases) - 1 and elapsed >= phases[phase_idx][0]:
                phase_idx += 1
                phase_label = phases[phase_idx][1]
                print(f"\r  {phase_label} {s}   ", end="", flush=True)
            print(f"\r  {phase_label} {s}   ", end="", flush=True)
        print(f"\r  Done ({elapsed:.0f}s).                                                      ")

        if "error" in _result:
            raise _result["error"]

        resp = _result["resp"]

        # The current Python SDK exposes convenience property output_text.
        text = getattr(resp, "output_text", None)
        if not text:
            # Fallback parser for SDK versions that do not expose output_text.
            chunks = []
            for item in getattr(resp, "output", []) or []:
                for content in getattr(item, "content", []) or []:
                    ctype = getattr(content, "type", "")
                    if ctype in {"output_text", "text"}:
                        chunks.append(getattr(content, "text", "") or "")
            text = "\n".join(chunks).strip()

        (debug_dir / "openai_analysis_response.md").write_text(text or "", encoding="utf-8")
        try:
            # Save a raw-ish JSON response when possible for troubleshooting citations/tool calls.
            if hasattr(resp, "model_dump"):
                (debug_dir / "openai_raw_response.json").write_text(json.dumps(resp.model_dump(), indent=2, default=str), encoding="utf-8")
            else:
                (debug_dir / "openai_raw_response.txt").write_text(str(resp), encoding="utf-8")
        except Exception:
            pass

        if not text:
            return None, f"OpenAI Responses API returned an empty response using model {model}; used local rules-based analysis."
        status = f"OpenAI web-search analysis generated using Responses API model {model}."
        return markdownish_to_html(text), status

    except Exception as e:
        err = f"{type(e).__name__}: {e}"
        (debug_dir / "openai_error.txt").write_text(err, encoding="utf-8")
        return None, f"OpenAI web-search analysis failed ({err}); used local rules-based analysis."


def rules_based_narrative(stock_agg: pd.DataFrame, all_agg: pd.DataFrame, buckets: pd.DataFrame, meta: Dict, prompt_text: str, positions: Optional[pd.DataFrame] = None) -> str:
    total = meta["total_value"]
    out = []

    def h2(t): out.append(f"<h2>{html.escape(t)}</h2>")
    def h3(t): out.append(f"<h3>{html.escape(t)}</h3>")
    def p(t):  out.append(f"<p>{t}</p>")
    def note(t): out.append(f"<p style='color:var(--muted);font-size:13px'>{t}</p>")

    # --- derived metrics ---
    def bucket_val(pattern: str) -> float:
        if buckets.empty: return 0.0
        return float(buckets[buckets["theme"].str.contains(pattern, case=False, na=False)]["value"].sum())

    crypto_val   = bucket_val("Crypto")
    ai_val       = bucket_val("AI|Semiconductors|Compute")
    biotech_val  = bucket_val("Biotech")
    energy_val   = bucket_val("Energy|Uranium|Nuclear")
    reit_val     = bucket_val("Real Estate|REIT")
    income_val   = bucket_val("Income|BDC")
    gold_val     = bucket_val("Gold|Precious")
    intl_val     = bucket_val("Emerging|International")
    broad_val    = bucket_val("Broad Market|S&P 500|Extended")
    bond_val     = bucket_val("Bond")
    cash_val     = bucket_val("Cash")
    space_val    = bucket_val("Space|Aviation")
    growth_val   = bucket_val("Large-Cap Growth")

    crypto_pct  = crypto_val / total * 100
    ai_pct      = ai_val / total * 100
    broad_pct   = broad_val / total * 100
    cash_pct    = cash_val / total * 100
    bond_pct    = bond_val / total * 100

    # concentration: top-5 and top-10 true stocks
    top5_pct  = float(stock_agg.head(5)["pct_portfolio"].sum())  if not stock_agg.empty else 0.0
    top10_pct = float(stock_agg.head(10)["pct_portfolio"].sum()) if not stock_agg.empty else 0.0
    over2  = stock_agg[stock_agg["pct_portfolio"] > 2.0]  if not stock_agg.empty else pd.DataFrame()
    over5  = stock_agg[stock_agg["pct_portfolio"] > 5.0]  if not stock_agg.empty else pd.DataFrame()

    # asset type mix from positions
    pos = positions.copy() if positions is not None and not positions.empty else pd.DataFrame()

    # -----------------------------------------------------------------------
    # SECTION 1 — Portfolio Snapshot
    # -----------------------------------------------------------------------
    h2("1. Portfolio Snapshot")
    out.append("<div style='display:grid;grid-template-columns:repeat(auto-fill,minmax(200px,1fr));gap:12px;margin:14px 0'>")
    def kpi(label, val, sub=""):
        sub_html = f"<div style='color:var(--muted);font-size:12px'>{sub}</div>" if sub else ""
        out.append(f"<div class='card'><div class='kpilabel'>{label}</div><div class='kpi' style='font-size:20px'>{val}</div>{sub_html}</div>")
    kpi("Total Portfolio", fmt_money(total))
    kpi("Crypto / Tokenization", fmt_money(crypto_val), f"{crypto_pct:.1f}% of portfolio")
    kpi("AI / Semis / Compute", fmt_money(ai_val), f"{ai_pct:.1f}% of portfolio")
    kpi("Broad Market Core", fmt_money(broad_val), f"{broad_pct:.1f}% of portfolio")
    kpi("Cash & Money Market", fmt_money(cash_val), f"{cash_pct:.1f}% of portfolio")
    kpi("Fixed Income / Bonds", fmt_money(bond_val), f"{bond_pct:.1f}% of portfolio")
    kpi("Top-5 Stock Concentration", f"{top5_pct:.1f}%", "of portfolio in 5 names")
    kpi("Top-10 Stock Concentration", f"{top10_pct:.1f}%", "of portfolio in 10 names")
    out.append("</div>")

    # -----------------------------------------------------------------------
    # SECTION 2 — Theme Breakdown & Observations
    # -----------------------------------------------------------------------
    h2("2. Theme Breakdown & Observations")
    if not buckets.empty:
        out.append("<table><thead><tr><th>#</th><th>Theme</th><th>Value</th><th>% Portfolio</th><th>Observation</th></tr></thead><tbody>")
        for i, (_, r) in enumerate(buckets.head(20).iterrows(), 1):
            v   = float(r["value"])
            pct = float(r["pct_portfolio"])
            theme = str(r["theme"])
            # generate a brief observation per theme
            if "Crypto" in theme:
                obs = "High-volatility sleeve. Watch against 10-15% guardrail." if pct > 10 else "Within guardrail band. Monitor for run-ups."
            elif "AI" in theme or "Semiconductor" in theme:
                obs = f"Core growth thesis. Largest mapped sleeve at {pct:.1f}%." if pct == ai_pct else "Part of AI/compute exposure."
            elif "Broad Market" in theme or "S&P 500" in theme:
                obs = "Core diversified base. Good anchor for volatile sleeves."
            elif "Bond" in theme:
                obs = "Defensive ballast. Low relative to growth tilt." if pct < 10 else "Meaningful fixed-income allocation."
            elif "Cash" in theme:
                obs = "Dry powder / emergency buffer." if pct < 5 else "High cash — consider deploying into underweight areas."
            elif "Biotech" in theme:
                obs = "High-risk / high-reward. Binary event exposure. Size carefully."
            elif "Uranium" in theme or "Nuclear" in theme:
                obs = "Long-cycle energy transition thesis. Illiquid names — watch position sizes."
            elif "Emerging" in theme:
                obs = "Geographic diversification. Currency and political risk embedded."
            elif "Real Estate" in theme or "REIT" in theme:
                obs = "Income + inflation hedge. Rate-sensitive."
            elif "Income" in theme or "BDC" in theme:
                obs = "Yield-focused. Credit risk embedded. Watch if rates stay high."
            elif "Gold" in theme or "Precious" in theme:
                obs = "Inflation/tail-risk hedge."
            elif "Space" in theme or "Aviation" in theme:
                obs = "Early-stage / pre-revenue exposure. Speculative sizing appropriate."
            elif "Fund Remainder" in theme or "OTHER" in theme:
                obs = "Un-mapped fund holdings. True exposure is broader than shown."
            else:
                obs = ""
            out.append(f"<tr><td>{i}</td><td>{html.escape(theme)}</td><td>{fmt_money(v)}</td><td>{fmt_pct(pct)}</td><td style='color:var(--muted);font-size:12px'>{html.escape(obs)}</td></tr>")
        out.append("</tbody></table>")

    # -----------------------------------------------------------------------
    # SECTION 3 — Concentration Risk
    # -----------------------------------------------------------------------
    h2("3. Concentration Risk")
    p(f"Your top 5 true stock exposures account for <b>{top5_pct:.1f}%</b> of the portfolio "
      f"and your top 10 account for <b>{top10_pct:.1f}%</b>. "
      + ("<b>Concentration is elevated</b> — a single bad earnings report or sector rotation can move the portfolio meaningfully." if top10_pct > 30 else
         "Concentration is moderate. The broad-market fund base helps buffer single-stock swings."))

    if not stock_agg.empty:
        out.append("<table><thead><tr><th>#</th><th>Ticker</th><th>Name</th><th>Value</th><th>% Portfolio</th><th>Flag</th></tr></thead><tbody>")
        for i, (_, r) in enumerate(stock_agg.head(20).iterrows(), 1):
            pct = float(r["pct_portfolio"])
            flag = ""
            if pct > 5:   flag = "<span style='color:#FF7A7A'>&#9650; High concentration (&gt;5%)</span>"
            elif pct > 2: flag = "<span style='color:var(--warn)'>&#9651; Watch (&gt;2%)</span>"
            out.append(f"<tr><td>{i}</td><td><b>{html.escape(str(r['exposure_ticker']))}</b></td>"
                       f"<td>{html.escape(str(r['exposure_name']))}</td>"
                       f"<td>{fmt_money(float(r['value']))}</td><td>{fmt_pct(pct)}</td><td>{flag}</td></tr>")
        out.append("</tbody></table>")

    # -----------------------------------------------------------------------
    # SECTION 4 — Guardrail Review
    # -----------------------------------------------------------------------
    h2("4. Guardrail Review")

    def guardrail_row(label, current_pct, target_lo, target_hi, invert=False):
        inside = target_lo <= current_pct <= target_hi
        if invert: inside = current_pct <= target_hi
        color  = "var(--good)" if inside else "var(--warn)"
        status = "PASS" if inside else "WATCH"
        out.append(f"<tr><td>{html.escape(label)}</td><td>{current_pct:.1f}%</td>"
                   f"<td>{target_lo:.0f}% – {target_hi:.0f}%</td>"
                   f"<td style='color:{color};font-weight:700'>{status}</td></tr>")

    out.append("<table><thead><tr><th>Guardrail</th><th>Current</th><th>Target Band</th><th>Status</th></tr></thead><tbody>")
    guardrail_row("Crypto / Tokenization",              crypto_pct,  0,  15)
    guardrail_row("AI / Semis / Compute",               ai_pct,      0,  25)
    guardrail_row("Broad market core (funds)",          broad_pct,   40, 100)
    guardrail_row("Cash & money market",                cash_pct,    0,  10)
    guardrail_row("Speculative stocks > 2% each",       len(over2),  0,   5, invert=True)
    guardrail_row("Speculative stocks > 5% each",       len(over5),  0,   1, invert=True)
    guardrail_row("Top-10 true stock concentration",    top10_pct,   0,  35)
    out.append("</tbody></table>")
    note("Target bands are illustrative defaults. Adjust them to match your own risk guardrails.")

    # -----------------------------------------------------------------------
    # SECTION 5 — Diversification Assessment
    # -----------------------------------------------------------------------
    h2("5. Diversification Assessment")
    themes_present = set(buckets["theme"].tolist()) if not buckets.empty else set()
    has_intl    = intl_val > 0
    has_bonds   = bond_val > 0
    has_reits   = reit_val > 0
    has_gold    = gold_val > 0
    has_income  = income_val > 0
    has_broad   = broad_val > 0

    out.append("<table><thead><tr><th>Dimension</th><th>Present?</th><th>Value</th><th>Note</th></tr></thead><tbody>")
    def div_row(label, present, val, note_txt):
        icon  = "&#10003;" if present else "&#8212;"
        color = "var(--good)" if present else "var(--muted)"
        out.append(f"<tr><td>{html.escape(label)}</td><td style='color:{color}'>{icon}</td>"
                   f"<td>{fmt_money(val) if val else '—'}</td>"
                   f"<td style='color:var(--muted);font-size:12px'>{html.escape(note_txt)}</td></tr>")
    div_row("Broad US market (index funds)", has_broad,  broad_val,  "Core anchor — good to have above 40%")
    div_row("International / Emerging",      has_intl,   intl_val,   "Geographic hedge against US-specific risk")
    div_row("Fixed income / Bonds",          has_bonds,  bond_val,   "Defensive ballast; reduces drawdown depth")
    div_row("Real Estate / REITs",           has_reits,  reit_val,   "Inflation hedge + income")
    div_row("Gold / Precious metals",        has_gold,   gold_val,   "Tail-risk / inflation hedge")
    div_row("Income / BDCs",                 has_income, income_val, "Yield-focused; credit risk")
    div_row("Crypto / Tokenization",         crypto_val > 0, crypto_val, "High volatility — size deliberately")
    div_row("AI / Growth theme",             ai_val > 0, ai_val,     "Concentrated growth bet")
    div_row("Biotech / Life Sciences",       biotech_val > 0, biotech_val, "Binary risk; high upside potential")
    div_row("Energy / Uranium / Nuclear",    energy_val > 0, energy_val, "Long-cycle transition thesis")
    out.append("</tbody></table>")

    missing = [lbl for lbl, present in [
        ("International / Emerging Markets", has_intl),
        ("Fixed Income / Bonds", has_bonds),
        ("Real Estate / REITs", has_reits),
    ] if not present]
    if missing:
        p(f"<b>Gaps to consider:</b> {', '.join(missing)} are not represented. These are not urgent additions — but they are the most common diversification levers if you want to reduce correlation to US growth equities.")

    # -----------------------------------------------------------------------
    # SECTION 6 — Speculative & High-Risk Positions
    # -----------------------------------------------------------------------
    h2("6. Speculative & High-Risk Position Review")
    spec_themes = ["Biotech", "Space", "Crypto", "Uranium", "AI Drug", "Gene Editing"]
    spec_rows = stock_agg[stock_agg["theme"].str.contains("|".join(spec_themes), case=False, na=False)] if not stock_agg.empty and "theme" in stock_agg.columns else pd.DataFrame()

    if not over2.empty:
        p("The following true stock exposures exceed 2% individually. That is not a sell signal — "
          "it is a sizing awareness flag. The question is whether the position size matches your conviction and risk tolerance.")
        out.append("<table><thead><tr><th>Ticker</th><th>Name</th><th>% Portfolio</th><th>Value</th><th>Action guidance</th></tr></thead><tbody>")
        for _, r in over2.iterrows():
            pct = float(r["pct_portfolio"])
            action = ("Consider trimming or pausing new contributions to this name." if pct > 5
                      else "Hold if thesis intact. Use new contributions elsewhere first.")
            out.append(f"<tr><td><b>{html.escape(str(r['exposure_ticker']))}</b></td>"
                       f"<td>{html.escape(str(r['exposure_name']))}</td>"
                       f"<td>{fmt_pct(pct)}</td><td>{fmt_money(float(r['value']))}</td>"
                       f"<td style='font-size:12px;color:var(--muted)'>{html.escape(action)}</td></tr>")
        out.append("</tbody></table>")
    else:
        p("No single true stock exposure exceeds 2%. Concentration risk is well-distributed at the individual name level.")

    # -----------------------------------------------------------------------
    # SECTION 7 — Rebalancing Guidance (contribution-first)
    # -----------------------------------------------------------------------
    h2("7. Rebalancing Guidance")
    p("Guidance below is contribution-first: redirect new money before trimming existing winners, "
      "unless a sleeve is clearly oversized or the original thesis has changed.")

    out.append("<table><thead><tr><th>Area</th><th>Current %</th><th>Guidance</th><th>Rationale</th></tr></thead><tbody>")

    def rebal_row(area, pct, guidance, rationale):
        color = {"Hold": "var(--good)", "Watch": "var(--warn)", "Increase contributions": "#7EC8E3",
                 "Reduce new contributions": "var(--warn)", "Consider trimming": "#FF7A7A"}.get(guidance, "var(--muted)")
        out.append(f"<tr><td>{html.escape(area)}</td><td>{pct:.1f}%</td>"
                   f"<td style='color:{color};font-weight:600'>{html.escape(guidance)}</td>"
                   f"<td style='font-size:12px;color:var(--muted)'>{html.escape(rationale)}</td></tr>")

    rebal_row("Broad market / core funds", broad_pct,
              "Increase contributions" if broad_pct < 40 else "Hold",
              "Core anchor should be the foundation before sizing satellite sleeves")
    rebal_row("Crypto / Tokenization", crypto_pct,
              "Consider trimming" if crypto_pct > 15 else ("Hold" if crypto_pct > 5 else "Watch"),
              "Above 15% guardrail — use profits to rebalance" if crypto_pct > 15 else "Within range; let winners run unless approaching guardrail")
    rebal_row("AI / Semis / Compute", ai_pct,
              "Hold" if ai_pct < 20 else "Reduce new contributions",
              "Thesis intact for long cycle; avoid adding aggressively if already oversized")
    rebal_row("Fixed income / Bonds", bond_pct,
              "Increase contributions" if bond_pct < 5 else "Hold",
              "Under-allocated to defensive ballast for a portfolio this growth-tilted")
    rebal_row("Cash", cash_pct,
              "Deploy into underweight areas" if cash_pct > 8 else "Hold",
              "High cash drag in a rising market unless intentionally waiting for entry")
    rebal_row("International / Emerging", intl_val / total * 100 if total else 0,
              "Increase contributions" if intl_val / total * 100 < 5 else "Hold",
              "Geographic diversification reduces US-only concentration risk")
    out.append("</tbody></table>")

    # -----------------------------------------------------------------------
    # SECTION 8 — Quarterly Decision Table
    # -----------------------------------------------------------------------
    h2("8. Quarterly Decision Table")
    note("A simple checklist for this review cycle. Not financial advice — use as a starting point.")
    out.append("<table><thead><tr><th>Decision</th><th>Action</th><th>Trigger to revisit</th></tr></thead><tbody>")

    decisions = [
        ("Crypto guardrail", "PASS — hold" if crypto_pct <= 15 else "WATCH — consider trimming",
         "If crypto sleeve exceeds 15% of total portfolio"),
        ("Broad core allocation", "Increase new contributions" if broad_pct < 40 else "Hold current allocation",
         "If broad core falls below 35% due to growth sleeve outperforming"),
        ("Speculative stocks > 2%", f"{len(over2)} positions flagged — review each thesis individually" if not over2.empty else "None — all clear",
         "Any single name approaches 5% of portfolio"),
        ("AI / compute sizing", "Hold — thesis intact" if ai_pct < 25 else "Pause new contributions",
         "If AI/semis bucket exceeds 25% or thesis shows cracks"),
        ("Fixed income allocation", "Consider adding" if bond_pct < 5 else "Hold",
         "Rising recession probability or rate cut cycle begins"),
        ("International exposure", "Consider adding via EM or international ETF" if not has_intl else "Hold",
         "USD weakens or EM outperformance cycle begins"),
        ("Fund look-through quality", "Add holdings_cache.csv for better mapping" if any("OTHER" in str(t) for t in themes_present) else "Look-through quality is good",
         "If fund OTHER bucket exceeds 20% of portfolio"),
    ]
    for decision, action, trigger in decisions:
        out.append(f"<tr><td>{html.escape(decision)}</td><td><b>{html.escape(action)}</b></td>"
                   f"<td style='font-size:12px;color:var(--muted)'>{html.escape(trigger)}</td></tr>")
    out.append("</tbody></table>")

    # -----------------------------------------------------------------------
    # SECTION 9 — Plain-English Bottom Line
    # -----------------------------------------------------------------------
    h2("9. Plain-English Bottom Line")

    bullets = []
    bullets.append(f"Total portfolio value: <b>{fmt_money(total)}</b>.")
    if crypto_pct > 15:
        bullets.append(f"Crypto is <b>{crypto_pct:.1f}%</b> — above the 15% guardrail. Trim or redirect contributions.")
    elif crypto_pct > 5:
        bullets.append(f"Crypto is <b>{crypto_pct:.1f}%</b> — inside the guardrail band. Hold unless it runs up.")
    else:
        bullets.append(f"Crypto is <b>{crypto_pct:.1f}%</b> — small allocation. Consider whether that matches conviction.")

    if broad_pct < 35:
        bullets.append(f"Broad market core is only <b>{broad_pct:.1f}%</b>. Your satellite sleeves (AI, crypto, biotech) are carrying a lot of the weight. Direct new contributions toward index funds first.")
    else:
        bullets.append(f"Broad market core is <b>{broad_pct:.1f}%</b> — a solid anchor for the growth sleeves.")

    if not over5.empty:
        names = ", ".join(str(r["exposure_ticker"]) for _, r in over5.iterrows())
        bullets.append(f"<b>{names}</b> each exceed 5% individually. Revisit whether position size still matches thesis.")

    if not has_bonds:
        bullets.append("No meaningful fixed income. This is an aggressive growth portfolio — make sure that matches your time horizon.")

    if not has_intl:
        bullets.append("No international / emerging markets exposure. All growth relies on US market outcomes.")

    bullets.append("Use new contributions to rebalance before trimming existing winners — selling early is usually the bigger long-term mistake.")

    out.append("<ul>" + "".join(f"<li>{b}</li>" for b in bullets) + "</ul>")

    if not prompt_text.strip():
        note("No quarterly review prompt found. Add a prompt.docx to the project folder to get a more personalized analysis and to enable the AI analysis tab.")

    return "\n".join(out)


def generate_html(positions: pd.DataFrame, all_agg: pd.DataFrame, stock_agg: pd.DataFrame, detail: pd.DataFrame, buckets: pd.DataFrame, meta: Dict, prompt_text: str, out: Path, inputs: Dict[str, Optional[Path]], use_openai: bool = True) -> None:
    total = meta["total_value"]
    positions = positions.copy()
    positions["pct_portfolio"] = positions["value"] / total * 100
    positions = positions.sort_values("value", ascending=False)

    stock50 = stock_agg.head(50).copy()
    stock50.insert(0, "rank", range(1, len(stock50) + 1))
    stock50 = append_total_row(stock50, "exposure_ticker", total, ["value"])

    all50 = all_agg.head(70).copy()
    all50.insert(0, "rank", range(1, len(all50) + 1))
    all50 = append_total_row(all50, "exposure_ticker", total, ["value"])

    bucket_table = buckets.copy()
    bucket_table.insert(0, "rank", range(1, len(bucket_table) + 1))
    bucket_table = append_total_row(bucket_table, "theme", total, ["value"])

    coverage = pd.DataFrame(meta.get("fund_coverage", []))
    if not coverage.empty:
        coverage["looked_through_value"] = coverage["fund_value"] * coverage["looked_through_pct"] / 100
        coverage["remainder_value"] = coverage["fund_value"] * coverage["remainder_pct"] / 100
        coverage = coverage.sort_values("fund_value", ascending=False)

    detail_view = detail[["exposure_ticker", "exposure_name", "source", "source_name", "source_weight_pct", "value", "pct_portfolio", "theme"]].copy()
    detail_view = detail_view.sort_values(["exposure_ticker", "value"], ascending=[True, False])

    # Reconciliation figures for each table
    stock_shown_val  = float(stock_agg.head(50)["value"].sum())
    stock_total_val  = float(stock_agg["value"].sum())
    stock_not_equity = total - stock_total_val  # fund remainders, crypto, pension, cash etc
    stock_off_bottom = stock_total_val - stock_shown_val

    all_shown_val    = float(all_agg.head(70)["value"].sum())
    all_off_bottom   = total - all_shown_val

    bucket_total_val = float(buckets["value"].sum())

    def recon(shown, grand_total, rows_shown, rows_total, extra_note=""):
        gap = grand_total - shown
        pct = gap / grand_total * 100 if grand_total else 0
        lines = [f"<b>Showing {rows_shown} of {rows_total} rows.</b> Shown total: <b>{fmt_money(shown)}</b> &nbsp;|&nbsp; Portfolio total: <b>{fmt_money(grand_total)}</b>"]
        if gap > 1:
            lines.append(f"&nbsp;|&nbsp; <span style='color:var(--warn)'>Gap: {fmt_money(gap)} ({pct:.1f}%)</span>")
        if extra_note:
            lines.append(f"<br><span style='color:var(--muted)'>{extra_note}</span>")
        return f"<div class='recon-box'>{''.join(lines)}</div>"

    stock_table = df_to_html_table(stock50, ["rank", "exposure_ticker", "exposure_name", "value", "pct_portfolio"], money_cols=["value"], pct_cols=["pct_portfolio"], table_id="stock50")
    stock_recon = recon(stock_shown_val, total, min(50, len(stock_agg)), len(stock_agg),
        f"This table shows <b>true equity look-through only</b>. The {fmt_money(stock_not_equity)} gap to portfolio total is fund remainders, crypto, pension/retirement, cash, and other non-equity buckets &mdash; not missing data.")

    all_table = df_to_html_table(all50, ["rank", "exposure_ticker", "exposure_name", "value", "pct_portfolio"], money_cols=["value"], pct_cols=["pct_portfolio"], table_id="all50")
    all_recon = recon(all_shown_val, total, min(70, len(all_agg)), len(all_agg),
        f"Showing top 70 of {len(all_agg)} exposures by value. The {fmt_money(all_off_bottom)} gap is smaller positions below the display cutoff. Theme Buckets above includes all {len(all_agg)} and reconciles to 100%." if all_off_bottom > 1 else "")

    bucket_html = df_to_html_table(bucket_table, ["rank", "theme", "value", "pct_portfolio"], money_cols=["value"], pct_cols=["pct_portfolio"], table_id="buckets")
    bucket_recon = recon(bucket_total_val, total, len(buckets), len(buckets),
        "Theme buckets include every position. This table ties to the portfolio total." if abs(bucket_total_val - total) < 1 else "")

    pos_html = df_to_html_table(positions, ["ticker", "name", "asset_type", "value", "pct_portfolio"], money_cols=["value"], pct_cols=["pct_portfolio"], table_id="positions")
    pos_recon = recon(float(positions["value"].sum()), total, len(positions), len(positions), "Original positions as parsed from your account files. This is the source of truth for the total.")

    cov_html = df_to_html_table(coverage, ["fund", "fund_name", "fund_value", "looked_through_pct", "looked_through_value", "remainder_pct", "remainder_value", "source"], money_cols=["fund_value", "looked_through_value", "remainder_value"], pct_cols=["looked_through_pct", "remainder_pct"], table_id="coverage") if not coverage.empty else "<p>No fund coverage records.</p>"
    detail_html = df_to_html_table(detail_view, ["exposure_ticker", "exposure_name", "source", "source_name", "source_weight_pct", "value", "pct_portfolio", "theme"], money_cols=["value"], pct_cols=["source_weight_pct", "pct_portfolio"], table_id="detail")

    # Always build the local analysis — it provides the structured KPI/table view.
    local_narrative = rules_based_narrative(stock_agg, all_agg, buckets, meta, prompt_text, positions=positions)

    # OpenAI adds the institutional outlook section on top of the local analysis.
    openai_narrative = None
    openai_status = ""
    if use_openai:
        openai_narrative, openai_status = generate_openai_narrative(positions, all_agg, stock_agg, detail, buckets, meta, prompt_text, Path(inputs.get("project_dir") or DEFAULT_PROJECT_DIR))
    elif not use_openai:
        openai_status = "--no-openai flag set."

    if openai_narrative:
        ai_block = f"<div style='border-top:2px solid var(--gold);margin-top:36px;padding-top:24px'><h2 style='margin-top:0'>Institutional Outlook &amp; Action Plan <span style='font-size:13px;font-weight:400;color:var(--muted)'>(OpenAI web search &mdash; {html.escape(openai_status)})</span></h2>{openai_narrative}</div>"
    else:
        ai_block = "" if not use_openai else f"<p class='note' style='margin-top:24px'>{html.escape(openai_status or 'OpenAI analysis not available.')}</p>"

    narrative = local_narrative + ai_block
    crypto_value = float(buckets[buckets["theme"].str.contains("Crypto", case=False, na=False)]["value"].sum()) if not buckets.empty else 0
    ai_value = float(buckets[buckets["theme"].str.contains("AI|Semiconductors|Compute", case=False, na=False)]["value"].sum()) if not buckets.empty else 0
    fund_remainder_value = float(detail[detail["is_remainder"] == True]["value"].sum()) if not detail.empty else 0

    max_bucket = float(buckets["value"].max()) if not buckets.empty else 1
    bars = []
    for _, r in buckets.head(18).iterrows():
        w = max(2, min(100, float(r["value"]) / max_bucket * 100))
        bars.append(f"<div class='barrow'><div class='barlabel'>{html.escape(str(r['theme']))}</div><div class='barwrap'><div class='bar' style='width:{w:.1f}%'></div></div><div class='barval'>{fmt_pct(float(r['pct_portfolio']))}</div></div>")

    input_notes = "<br>".join([
        f"Project folder: {html.escape(str(inputs.get('project_dir', '')))}",
        f"Positions: {html.escape(str(inputs.get('positions', '')))}",
        f"Prompt file: {html.escape(str(inputs.get('prompt_file') or inputs.get('prompt_docx') or ''))}",
        f"Holdings cache: {html.escape(str(inputs.get('holdings_cache') or 'none'))}",
        f"Output: {html.escape(str(out))}",
        f"Institutional outlook: {'OpenAI web search (' + openai_status + ')' if openai_narrative else ('Add OPENAI_API_KEY to .env to enable' if not use_openai or not openai_status else openai_status)}",
    ])

    css = """
    :root{--navy:#071426;--gold:#C9A84C;--ink:#EAF2FF;--muted:#A8B3C6;--card:#111F34;--line:#2B3C58;--good:#7EE2A8;--warn:#FFD37A;}
    *{box-sizing:border-box} body{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;background:linear-gradient(180deg,var(--navy),#050912);color:var(--ink)}
    header{padding:34px 42px;border-bottom:1px solid var(--line);background:radial-gradient(circle at 70% 0%,rgba(201,168,76,.20),transparent 38%)}
    h1{margin:0;color:var(--gold);font-size:34px;letter-spacing:.4px} h2{color:var(--gold);margin-top:26px} h3{color:#fff}.sub{color:var(--muted);margin-top:8px}.small{font-size:12px;color:var(--muted)}
    .tabs{display:flex;gap:10px;padding:18px 42px 0;flex-wrap:wrap}.tabbtn{border:1px solid var(--line);background:#0d1b2f;color:var(--ink);padding:12px 18px;border-radius:10px;cursor:pointer;font-weight:700}.tabbtn.active{background:var(--gold);color:#071426}
    .tab{display:none;padding:24px 42px 50px}.tab.active{display:block}.grid{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:16px}.card{background:rgba(17,31,52,.92);border:1px solid var(--line);border-radius:18px;padding:18px;box-shadow:0 10px 30px rgba(0,0,0,.25)}
    .kpi{font-size:26px;font-weight:800;color:#fff}.kpilabel{color:var(--muted);font-size:13px;text-transform:uppercase;letter-spacing:.08em}
    table{width:100%;border-collapse:collapse;background:rgba(17,31,52,.8);border:1px solid var(--line);border-radius:14px;overflow:hidden;margin:14px 0 24px}
    th,td{padding:10px 12px;border-bottom:1px solid var(--line);text-align:left;font-size:13px}
    th{color:var(--gold);font-size:12px;text-transform:uppercase;letter-spacing:.05em;background:#0b1728}
    th:hover{background:#162842} .sort-icon{opacity:.4;font-size:10px} th.asc .sort-icon::after{content:'\\25B2'} th.desc .sort-icon::after{content:'\\25BC'}
    tr:hover{background:#162842}
    tr.total-row td{font-weight:700;color:var(--gold);background:#0b1728;border-top:2px solid var(--gold)}
    .note{color:var(--muted);font-size:13px}.recon-box{background:#0b1728;border:1px solid var(--line);border-radius:10px;padding:12px 16px;margin:6px 0 16px;font-size:13px}
    .barrow{display:grid;grid-template-columns:280px 1fr 75px;align-items:center;gap:10px;margin:10px 0}.barlabel{color:#fff;font-size:13px}.barwrap{height:14px;background:#0b1728;border:1px solid var(--line);border-radius:999px;overflow:hidden}.bar{height:100%;background:linear-gradient(90deg,var(--gold),#fff2b8)}.barval{font-size:13px;color:var(--muted);text-align:right}.footer{padding:20px 42px;color:var(--muted);border-top:1px solid var(--line)}
    input.search{width:100%;padding:12px 14px;background:#0b1728;border:1px solid var(--line);border-radius:12px;color:var(--ink);margin:10px 0 12px}.warn{color:var(--warn)}
    @media(max-width:900px){.grid{grid-template-columns:1fr}.barrow{grid-template-columns:1fr}.tab,header,.tabs,.footer{padding-left:18px;padding-right:18px}}
    """
    js = """
    function showTab(id){document.querySelectorAll('.tab').forEach(x=>x.classList.remove('active'));document.querySelectorAll('.tabbtn').forEach(x=>x.classList.remove('active'));document.getElementById(id).classList.add('active');document.getElementById('btn_'+id).classList.add('active');}
    function filterTable(inputId, tableId){const q=document.getElementById(inputId).value.toLowerCase();document.querySelectorAll('#'+tableId+' tbody tr').forEach(tr=>{if(tr.classList.contains('total-row'))return;tr.style.display=tr.innerText.toLowerCase().includes(q)?'':'none';});}
    function sortTable(tableId, th){
      const tbl=document.getElementById(tableId);
      const tbody=tbl.querySelector('tbody');
      const col=parseInt(th.dataset.col);
      const asc=th.classList.contains('asc');
      tbl.querySelectorAll('thead th').forEach(h=>h.classList.remove('asc','desc'));
      th.classList.add(asc?'desc':'asc');
      const rows=[...tbody.querySelectorAll('tr:not(.total-row)')];
      const totals=[...tbody.querySelectorAll('tr.total-row')];
      rows.sort((a,b)=>{
        const av=a.cells[col]?.dataset?.val??a.cells[col]?.innerText??'';
        const bv=b.cells[col]?.dataset?.val??b.cells[col]?.innerText??'';
        const an=parseFloat(av),bn=parseFloat(bv);
        if(!isNaN(an)&&!isNaN(bn))return asc?bn-an:an-bn;
        return asc?bv.localeCompare(av):av.localeCompare(bv);
      });
      rows.forEach(r=>tbody.appendChild(r));
      totals.forEach(r=>tbody.appendChild(r));
    }
    """

    html_doc = f"""<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Portfolio X-Ray Report</title><style>{css}</style></head>
<body><header><h1>Portfolio X-Ray & Quarterly Rotation Review</h1><div class='sub'>Generated {html.escape(meta['generated_at'])}. Top {TOP_HOLDINGS_PER_FUND} fund holdings are x-rayed where available; the rest of each fund remains in a clear OTHER bucket. Not financial advice.</div><div class='small'>{input_notes}</div></header>
<div class='tabs'><button id='btn_xray' class='tabbtn active' onclick="showTab('xray')">Portfolio X-Ray</button><button id='btn_review' class='tabbtn' onclick="showTab('review')">Rotation Review</button></div>
<section id='xray' class='tab active'>
<div class='grid'><div class='card'><div class='kpilabel'>Total analyzed value</div><div class='kpi'>{fmt_money(total)}</div></div><div class='card'><div class='kpilabel'>Crypto/tokenization</div><div class='kpi'>{fmt_money(crypto_value)}</div><div class='note'>{crypto_value/total*100:.2f}% of portfolio</div></div><div class='card'><div class='kpilabel'>AI / compute mapped</div><div class='kpi'>{fmt_money(ai_value)}</div><div class='note'>{ai_value/total*100:.2f}% of portfolio</div></div><div class='card'><div class='kpilabel'>Fund OTHER / unmapped</div><div class='kpi'>{fmt_money(fund_remainder_value)}</div><div class='note'>{fund_remainder_value/total*100:.2f}% of portfolio</div></div></div>
<h2>Exposure Buckets</h2><p class='note'>Every position is in exactly one bucket. This table always ties to 100% of your portfolio. Click any column header to sort.</p><div class='card'>{''.join(bars)}</div>{bucket_recon}{bucket_html}
<h2>Top 50 True Stock Concentrations</h2><p class='note'>Direct stocks plus look-through holdings inside funds, aggregated by company. Excludes fund remainders, crypto, pension, and cash — those appear in the All Exposures table. Click any column to sort.</p><input class='search' id='s_stock' onkeyup="filterTable('s_stock','stock50')" placeholder='Search stocks...'>{stock_recon}{stock_table}
<h2>All Exposures Including Fund OTHER Buckets</h2><p class='note'>Full reconciliation view. Includes fund remainders so nothing is hidden. Click any column to sort.</p><input class='search' id='s_all' onkeyup="filterTable('s_all','all50')" placeholder='Search all exposures...'>{all_recon}{all_table}
<h2>Original Positions from Account Files</h2><p class='note'>Raw positions as parsed — one row per holding per account file. This is the source of truth for the portfolio total.</p><input class='search' id='s_pos' onkeyup="filterTable('s_pos','positions')" placeholder='Search positions...'>{pos_recon}{pos_html}
<h2>Fund Look-Through Coverage</h2>{cov_html}
<h2>Underlying Exposure Detail by Source</h2><p class='note'>Use this to answer questions like: how much Apple came from Contrafund versus FXAIX?</p><input class='search' id='s_detail' onkeyup="filterTable('s_detail','detail')" placeholder='Search source detail...'>{detail_html}
</section>
<section id='review' class='tab'><div class='card'><p class='note'>{html.escape(openai_status)}</p>{narrative}</div><h2>Prompt Used as Guide</h2><p class='note'>The narrative follows the goals/guardrails extracted from your DOCX prompt, including contribution-first rebalancing, crypto guardrails, and single-speculative-stock sizing awareness.</p><details><summary>Show extracted prompt text</summary><pre style='white-space:pre-wrap;color:#A8B3C6'>{html.escape(prompt_text[:15000])}</pre></details></section>
<div class='footer'>Method: local parsing only; holdings_cache.csv first, then yfinance if installed/available. Mutual fund holdings are delayed/incomplete, so verify before acting.</div><script>{js}</script></body></html>"""
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(html_doc, encoding="utf-8")


def save_history(project_dir: Path, positions: pd.DataFrame, all_agg: pd.DataFrame, stock_agg: pd.DataFrame, buckets: pd.DataFrame, meta: Dict) -> None:
    hist_dir = project_dir / "history"
    hist_dir.mkdir(parents=True, exist_ok=True)
    stamp = dt.datetime.now().strftime("%Y_%m_%d_%H%M")
    payload = {
        "generated_at": meta["generated_at"],
        "total_value": meta["total_value"],
        "top_stocks": stock_agg.head(20)[["exposure_ticker", "exposure_name", "value", "pct_portfolio"]].to_dict(orient="records"),
        "top_all_exposures": all_agg.head(20)[["exposure_ticker", "exposure_name", "value", "pct_portfolio"]].to_dict(orient="records"),
        "buckets": buckets[["theme", "value", "pct_portfolio"]].to_dict(orient="records"),
        "positions": positions[["ticker", "name", "value", "category"]].to_dict(orient="records"),
    }
    (hist_dir / f"portfolio_xray_{stamp}.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Build Will's Portfolio X-Ray HTML report.")
    ap.add_argument("--project-dir", help="Finance dashboard folder. Default: C:\\Users\\willb\\myapps\\finance-dashboard")
    ap.add_argument("--positions", help="Optional positions PDF/CSV/XLSX. If omitted, newest matching file in project dir is used.")
    ap.add_argument("--prompt-docx", help="Optional quarterly prompt DOCX/TXT. If omitted, newest matching prompt file in project dir is used.")
    ap.add_argument("--prompt-file", help="Optional quarterly prompt DOCX/TXT. Alias that supports .txt directly.")
    ap.add_argument("--holdings-cache", help="Optional holdings_cache.csv. If omitted, project-dir/holdings_cache.csv is used when present.")
    ap.add_argument("--out", help="Optional output HTML path. If omitted, writes to project-dir/output/portfolio_xray_TIMESTAMP.html and latest_portfolio_xray.html.")
    ap.add_argument("--no-history", action="store_true", help="Do not write a history JSON snapshot.")
    ap.add_argument("--no-openai", action="store_true", help="Disable OpenAI analysis tab generation even if OPENAI_API_KEY is present.")
    args = ap.parse_args(argv)

    project_dir = resolve_project_dir(args.project_dir)
    inputs = discover_inputs(project_dir, args.positions, args.prompt_file or args.prompt_docx, args.holdings_cache, args.out)
    inputs["project_dir"] = project_dir

    log(f"Project folder: {project_dir}")
    log(f"Prompt file: {inputs.get('prompt_file') if inputs.get('prompt_file') and inputs.get('prompt_file').exists() else 'not found'}")
    if inputs.get('prompt_file') and str(inputs.get('prompt_file')).lower().find('requirements') >= 0:
        raise SystemExit('Refusing to use a requirements file as the prompt. Rename your review prompt to include quarterly/rotation/review/prompt, or pass --prompt-file explicitly.')
    log(f"Holdings cache: {inputs['holdings_cache'] if inputs['holdings_cache'] and inputs['holdings_cache'].exists() else 'not found; will try yfinance'}")

    if args.positions:
        # Explicit override: load just that one file
        pos_path = inputs["positions"]
        if not pos_path or not pos_path.exists():
            raise SystemExit(f"Positions file not found: {pos_path}")
        log(f"Positions file (explicit): {pos_path}")
        positions, total_value = read_positions_table(pos_path)
        loaded_files = [pos_path]
    else:
        log("Loading all asset files...")
        positions, total_value, loaded_files = read_all_asset_files(project_dir)
        if positions.empty:
            raise SystemExit(f"No positions found. Put PDF/XLSX/CSV/JSON/TXT files in: {project_dir / 'assets'}")

    inputs["positions"] = ", ".join(p.name for p in loaded_files)

    if positions.empty:
        raise SystemExit("No positions parsed.")
    prompt_text = read_prompt_file(inputs.get("prompt_file") or inputs.get("prompt_docx"))
    cache = read_holdings_cache(inputs["holdings_cache"])
    all_agg, stock_agg, detail, buckets, meta = build_xray(positions, cache, total_value)

    out_path: Path = inputs["out"]
    generate_html(positions, all_agg, stock_agg, detail, buckets, meta, prompt_text, out_path, inputs, use_openai=not args.no_openai)

    if not args.no_history:
        save_history(project_dir, positions, all_agg, stock_agg, buckets, meta)

    log(f"Wrote: {out_path}")
    log(f"Parsed {len(positions)} positions from {len(loaded_files)} file(s), total value {fmt_money(meta['total_value'])}.")
    log("Top stock concentrations:")
    if stock_agg.empty:
        log("No stock look-through records available. Add holdings_cache.csv or ensure yfinance is installed and online.")
    else:
        print(stock_agg.head(10)[["exposure_ticker", "exposure_name", "value", "pct_portfolio"]].to_string(index=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
